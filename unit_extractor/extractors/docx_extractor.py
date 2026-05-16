from .base import ExtractorBase
from domain import TextData
from docx import Document



class DocxExtractor(ExtractorBase):
    """
    Extracts word document files.
    """
    def _extract(self, path=None) -> list:
        
        word_doc = Document(path)
        
        text = self._get_text(word_doc)
        table = self._get_table(word_doc)
            
        return [TextData(path=path, text=f'TEXT:{text}\nTABLES:{table}\n')]
    
    def _get_text(self, doc) -> str:
        """Concatenate non-empty paragraph text."""
        all_text = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                all_text.append(text)
                
        # join collected paragraphs with newline for readability
        return "\n".join(all_text)
    
    def _get_table(self, doc) -> str:
        """Extract table rows into tab-separated strings."""
        all_table_contents = []
        
        for table in doc.tables:
            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_data:
                    all_table_contents.append("\t".join(row_data))
        return "\n".join(all_table_contents)

if __name__ == "__main__":
    
    
    test_path = r"C:\All University Materials\Project\ICT304-project\ai\ai_tools\data_extractor\data\raw\ICT283_contents\Topic01\Topic1-Everything_extracted\Practical\Exc1 for revision\ICT283 Revision - ICT159 - Assignment2.docx"
    
    for i in DocxExtractor().extract(path=test_path):
        print(i.text)